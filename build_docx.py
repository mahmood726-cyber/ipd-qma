"""
Build a submission-ready PLOS ONE Word document from the corrected markdown draft.

Uses python-docx to create a properly formatted .docx with:
- PLOS ONE heading hierarchy
- Numbered tables with captions
- Figure legends
- Proper reference formatting
- All mandatory PLOS ONE sections
"""

import sys
import os
import re
import io

try:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
except Exception:
    pass

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DRAFT_PATH = os.path.join(SCRIPT_DIR, "IPD_QMA_corrected_draft.md")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "IPD_QMA_submission.docx")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output")

# Map draft figure numbers to output filenames
FIGURE_MAP = {
    '1': 'fig4_type_i_error.png',
    '2': 'fig5_power_by_vr.png',
    '3': 'fig6_power_by_k.png',
    '4': 'fig7_power_by_n.png',
    '5': 'fig1_nhanes_qte_profile.png',
    '6': 'fig2_nhanes_forest_plots.png',
    '7': 'fig3_nhanes_comparison.png',
    '8': 'fig8_coverage.png',
}

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def add_formatted_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def apply_body_style(paragraph, font_name='Times New Roman', font_size=12):
    """Apply standard body text style."""
    paragraph.style.font.name = font_name
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def add_table_from_lines(doc, header_line, data_lines, caption=None):
    """Build a Word table from markdown-style | delimited lines."""
    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    headers = parse_row(header_line)
    n_cols = len(headers)

    # Add caption above table
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_formatted_run(cap_para, caption, bold=True, size=10)
        cap_para.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Times New Roman'
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for line in data_lines:
        cells = parse_row(line)
        row = table.add_row()
        for i, val in enumerate(cells):
            if i < n_cols:
                cell = row.cells[i]
                cell.text = val
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = 'Times New Roman'

    # Add spacing after table
    doc.add_paragraph()
    return table


# ── Main builder ─────────────────────────────────────────────────────────────

def build_docx():
    print("Reading corrected draft...")
    with open(DRAFT_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Add continuous line numbering (PLOS ONE requirement)
        sectPr = section._sectPr
        lnNumType = sectPr.makeelement(qn('w:lnNumType'), {
            qn('w:countBy'): '1',
            qn('w:restart'): 'continuous',
        })
        sectPr.append(lnNumType)

    # ── Default styles ───────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0  # Double-spaced per PLOS ONE
    style.paragraph_format.space_after = Pt(0)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            hs.font.size = Pt(14)
            hs.font.bold = True
        elif level == 2:
            hs.font.size = Pt(12)
            hs.font.bold = True
        else:
            hs.font.size = Pt(12)
            hs.font.bold = True
            hs.font.italic = True

    # ── Page numbers in footer ─────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add PAGE field
        run = fp.add_run()
        fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._element.append(fldChar1)
        run2 = fp.add_run()
        instrText = run2._element.makeelement(qn('w:instrText'), {})
        instrText.text = ' PAGE '
        run2._element.append(instrText)
        run3 = fp.add_run()
        fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._element.append(fldChar2)

    # ── Title ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(
        title_para,
        "IPD-QMA: Detecting Heterogeneous Treatment Effects "
        "via Quantile Meta-Analysis of Individual Participant Data",
        bold=True, size=16
    )
    title_para.paragraph_format.space_after = Pt(6)

    # Author placeholder
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(author_para, "[Author names and affiliations to be completed]",
                      italic=True, size=11, color=RGBColor(128, 128, 128))
    author_para.paragraph_format.space_after = Pt(24)

    # ── Parse and render ─────────────────────────────────────────────────
    i = 0
    # Skip the markdown title line (first # heading)
    while i < len(lines) and not lines[i].startswith('## '):
        i += 1

    current_table_caption = None
    in_italic_block = False

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            # H4 -> bold italic paragraph (PLOS ONE only has 3 heading levels)
            para = doc.add_paragraph()
            add_formatted_run(para, text, bold=True, italic=True, size=11)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        elif stripped.startswith('### '):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        elif stripped.startswith('## '):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # ── Tables (markdown pipe format) ────────────────────────────
        if stripped.startswith('|') and '|' in stripped[1:]:
            # Collect caption: look back for **Table N.** line
            caption = current_table_caption
            current_table_caption = None

            header_line = stripped
            i += 1
            # Skip separator (|---|---|...)
            if i < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[i]):
                i += 1
            # Collect data rows
            data_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                data_lines.append(lines[i].strip())
                i += 1

            add_table_from_lines(doc, header_line, data_lines, caption)
            continue

        # ── Detect table captions for next table ─────────────────────
        if stripped.startswith('**Table '):
            current_table_caption = stripped.replace('**', '')
            i += 1
            continue

        # ── Equation labels (centered, italic) ───────────────────────
        if stripped.startswith('**Equation '):
            eq_label = stripped.replace('**', '')
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(para, eq_label, bold=True, italic=True, size=11)
            i += 1
            continue

        # ── Figure legend entries (with embedded image) ──────────────
        if stripped.startswith('**Figure '):
            # Extract figure number
            fig_match = re.match(r'\*\*Figure (\d+)', stripped)
            if fig_match:
                fig_num = fig_match.group(1)
                fig_file = FIGURE_MAP.get(fig_num)
                if fig_file:
                    fig_path = os.path.join(OUTPUT_DIR, fig_file)
                    if os.path.exists(fig_path):
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        run.add_picture(fig_path, width=Inches(5.5))
            # Add caption text
            para = doc.add_paragraph()
            _add_inline_formatted(para, stripped)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # ── Note/caveat lines ─────────────────────────────────────────
        if stripped.startswith('**Note'):
            para = doc.add_paragraph()
            _add_inline_formatted(para, stripped)
            i += 1
            continue

        # ── Numbered/bulleted lists ──────────────────────────────────
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            para = doc.add_paragraph(style='List Number')
            _add_inline_formatted(para, text)
            i += 1
            continue

        if stripped.startswith('- '):
            text = stripped[2:]
            para = doc.add_paragraph(style='List Bullet')
            _add_inline_formatted(para, text)
            i += 1
            continue

        # ── Blank lines ──────────────────────────────────────────────
        if stripped == '':
            i += 1
            continue

        # ── Italic block (starts and ends with *...*) ────────────────
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            para = doc.add_paragraph()
            text = stripped.strip('*').strip()
            add_formatted_run(para, text, italic=True, size=10)
            i += 1
            continue

        # ── Regular paragraphs ───────────────────────────────────────
        # Accumulate multi-line paragraphs
        para_text = stripped
        i += 1
        # Don't merge if next line is a heading, blank, table, or list
        while i < len(lines):
            next_line = lines[i].strip()
            if (next_line == '' or next_line.startswith('#') or
                next_line.startswith('|') or next_line.startswith('- ') or
                next_line.startswith('---') or re.match(r'^\d+\.\s', next_line) or
                next_line.startswith('**Table ') or next_line.startswith('**Figure ')):
                break
            para_text += ' ' + next_line
            i += 1

        para = doc.add_paragraph()
        _add_inline_formatted(para, para_text)

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")
    print(f"Size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


def _add_inline_formatted(paragraph, text):
    """Parse inline markdown (bold, italic, bold+italic) and add runs."""
    # Pattern: **bold**, *italic*, ***bold+italic***, [N] references
    # Process segments between formatting markers
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)

    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            inner = part[3:-3]
            add_formatted_run(paragraph, inner, bold=True, italic=True)
        elif part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            add_formatted_run(paragraph, inner, bold=True)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            inner = part[1:-1]
            add_formatted_run(paragraph, inner, italic=True)
        else:
            if part:
                paragraph.add_run(part)


# ── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    build_docx()
